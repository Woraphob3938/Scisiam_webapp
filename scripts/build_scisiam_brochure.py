from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Scisiam_app")
ASSETS = ROOT / "output" / "doc" / "assets"
OUTPUT = ROOT / "output" / "doc" / "แผ่นพับ_SciSiam_สำหรับกรรมการ.docx"

FONT = "TH Sarabun New"
NAVY = "0F172A"
BLUE = "2563EB"
BLUE_DARK = "1D4ED8"
PALE_BLUE = "DBEAFE"
VERY_PALE = "F8FBFF"
SLATE = "475569"
MUTED = "64748B"
LINE = "BFDBFE"
WHITE = "FFFFFF"
NOTE = "0C4A6E"


def hex_color(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, size: float, bold: bool = False, color: str = NAVY) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = hex_color(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    lang = run._element.get_or_add_rPr().find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.get_or_add_rPr().append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:eastAsia"), "th-TH")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_width(cell, width_cm: float) -> None:
    width = Cm(width_cm).twips
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, *, top=0.38, start=0.42, bottom=0.32, end=0.42) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(Cm(value).twips))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = LINE, size: int = 8) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def clear_cell(cell) -> None:
    cell._tc.clear_content()


def configure_paragraph(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_together = True


def add_text(
    cell,
    text: str,
    *,
    size: float = 16,
    bold: bool = False,
    color: str = NAVY,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
) :
    paragraph = cell.add_paragraph()
    configure_paragraph(paragraph, align=align, before=before, after=after, line=line)
    for index, part in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_label(cell, text: str, *, fill: str = BLUE, color: str = WHITE) -> None:
    table = cell.add_table(rows=1, cols=1)
    table.autofit = False
    set_no_borders(table)
    label_cell = table.cell(0, 0)
    set_cell_shading(label_cell, fill)
    set_cell_margins(label_cell, top=0.08, start=0.18, bottom=0.05, end=0.18)
    clear_cell(label_cell)
    add_text(
        label_cell,
        text,
        size=13.5,
        bold=True,
        color=color,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.0,
    )


def add_divider(cell, *, color: str = LINE, after: float = 4) -> None:
    paragraph = cell.add_paragraph()
    configure_paragraph(paragraph, after=after)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "9")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_picture(cell, path: Path, *, width_cm: float, after: float = 2) -> None:
    paragraph = cell.add_paragraph()
    configure_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, after=after)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_bullet(cell, text: str, *, size: float = 15.2, color: str = NAVY, after: float = 1.6) -> None:
    paragraph = cell.add_paragraph()
    configure_paragraph(paragraph, after=after, line=1.05)
    paragraph.paragraph_format.left_indent = Cm(0.22)
    paragraph.paragraph_format.first_line_indent = Cm(-0.22)
    mark = paragraph.add_run("●  ")
    set_run_font(mark, size=8, bold=True, color=BLUE)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color)


def add_check(cell, text: str, *, size: float = 15.0, after: float = 1.5) -> None:
    paragraph = cell.add_paragraph()
    configure_paragraph(paragraph, after=after, line=1.0)
    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.first_line_indent = Cm(-0.26)
    mark = paragraph.add_run("✓  ")
    set_run_font(mark, size=14, bold=True, color=BLUE)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=NAVY)


def add_step(cell, number: str, title: str, body: str) -> None:
    table = cell.add_table(rows=1, cols=2)
    table.autofit = False
    set_no_borders(table)
    number_cell, text_cell = table.rows[0].cells
    set_cell_width(number_cell, 0.9)
    set_cell_width(text_cell, 6.9)
    set_cell_shading(number_cell, BLUE)
    set_cell_margins(number_cell, top=0.12, start=0.05, bottom=0.08, end=0.05)
    set_cell_margins(text_cell, top=0.04, start=0.18, bottom=0.06, end=0.0)
    number_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    clear_cell(number_cell)
    clear_cell(text_cell)
    add_text(number_cell, number, size=17, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph = text_cell.add_paragraph()
    configure_paragraph(paragraph, after=1, line=1.0)
    run = paragraph.add_run(title)
    set_run_font(run, size=16, bold=True, color=NAVY)
    paragraph = text_cell.add_paragraph()
    configure_paragraph(paragraph, after=1, line=1.0)
    run = paragraph.add_run(body)
    set_run_font(run, size=13.5, color=SLATE)


def add_metrics(cell) -> None:
    table = cell.add_table(rows=1, cols=3)
    table.autofit = False
    set_no_borders(table)
    values = (("103", "ห้องทดลอง"), ("5", "กลุ่มวิชา"), ("2", "รูปแบบใช้งาน"))
    for target, (number, label) in zip(table.rows[0].cells, values):
        set_cell_shading(target, PALE_BLUE)
        set_cell_margins(target, top=0.15, start=0.05, bottom=0.12, end=0.05)
        clear_cell(target)
        add_text(target, number, size=23, bold=True, color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(target, label, size=12.5, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_subject(cell, thai: str, english: str) -> None:
    table = cell.add_table(rows=1, cols=2)
    table.autofit = False
    set_no_borders(table)
    dot, text_cell = table.rows[0].cells
    set_cell_width(dot, 0.55)
    set_cell_width(text_cell, 7.45)
    set_cell_shading(dot, BLUE)
    set_cell_shading(text_cell, WHITE)
    set_cell_margins(dot, top=0.13, start=0.02, bottom=0.10, end=0.02)
    set_cell_margins(text_cell, top=0.08, start=0.2, bottom=0.08, end=0.06)
    clear_cell(dot)
    clear_cell(text_cell)
    add_text(dot, "•", size=16, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph = text_cell.add_paragraph()
    configure_paragraph(paragraph, line=1.0)
    run = paragraph.add_run(thai)
    set_run_font(run, size=16.5, bold=True, color=NAVY)
    run = paragraph.add_run(f"  {english}")
    set_run_font(run, size=12.5, color=MUTED)


def add_info_card(cell, title: str, body: str, *, image: Path | None = None) -> None:
    table = cell.add_table(rows=1, cols=2 if image else 1)
    table.autofit = False
    set_table_borders(table, color=LINE, size=7)
    if image:
        image_cell, text_cell = table.rows[0].cells
        set_cell_width(image_cell, 2.0)
        set_cell_width(text_cell, 5.8)
        set_cell_shading(image_cell, WHITE)
        set_cell_shading(text_cell, WHITE)
        set_cell_margins(image_cell, top=0.12, start=0.10, bottom=0.10, end=0.10)
        set_cell_margins(text_cell, top=0.15, start=0.22, bottom=0.12, end=0.18)
        clear_cell(image_cell)
        clear_cell(text_cell)
        add_picture(image_cell, image, width_cm=1.65, after=0)
    else:
        text_cell = table.cell(0, 0)
        set_cell_shading(text_cell, WHITE)
        set_cell_margins(text_cell, top=0.15, start=0.22, bottom=0.12, end=0.18)
        clear_cell(text_cell)
    add_text(text_cell, title, size=17, bold=True, color=BLUE_DARK, after=1)
    add_text(text_cell, body, size=14, color=SLATE, line=1.05)


def add_note(cell, text: str) -> None:
    table = cell.add_table(rows=1, cols=1)
    table.autofit = False
    set_no_borders(table)
    note_cell = table.cell(0, 0)
    set_cell_shading(note_cell, PALE_BLUE)
    set_cell_margins(note_cell, top=0.13, start=0.18, bottom=0.10, end=0.18)
    clear_cell(note_cell)
    add_text(note_cell, text, size=13, color=NOTE, line=1.05)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(0.6)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.6)
    section.right_margin = Cm(0.6)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    # A 1 pt base keeps Word's mandatory structural paragraphs from creating
    # blank brochure pages; every visible run receives an explicit size below.
    normal.font.size = Pt(1)
    normal.font.color.rgb = hex_color(NAVY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(1)

    props = document.core_properties
    props.title = "แผ่นพับแนะนำ SciSiam Virtual Lab"
    props.subject = "แผ่นพับประชาสัมพันธ์สำหรับนำเสนอต่อกรรมการ"
    props.author = "SciSiam"
    props.keywords = "SciSiam, Virtual Lab, ห้องปฏิบัติการเสมือนจริง, วิทยาศาสตร์"


def configure_section(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(0.6)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.6)
    section.right_margin = Cm(0.6)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)


def create_panel_table(document: Document, fills: tuple[str, str, str]):
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_borders(table, color=LINE, size=8)
    row = table.rows[0]
    # Leave enough room for Word's mandatory paragraph after the table so
    # the two brochure sides remain exactly two pages when exported.
    row.height = Cm(19.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    widths = (9.3, 9.6, 9.6)
    for cell, width, fill in zip(row.cells, widths, fills):
        set_cell_width(cell, width)
        set_cell_shading(cell, fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        clear_cell(cell)
    return table, row.cells


def build_outside(document: Document) -> None:
    _, (flap, back, cover) = create_panel_table(document, (VERY_PALE, WHITE, PALE_BLUE))

    # Panel 5: fold-in flap / call to action
    add_label(flap, "เริ่มต้นใช้งานง่าย", fill=BLUE_DARK)
    add_text(flap, "พร้อมทดลอง\nใน 3 ขั้นตอน", size=25, bold=True, color=NAVY, before=5, after=5, line=0.92)
    add_step(flap, "1", "เปิดเว็บไซต์", "สแกนคิวอาร์โค้ด หรือเข้าสู่เว็บไซต์ SciSiam")
    add_step(flap, "2", "สมัครสมาชิก", "สร้างบัญชีผู้ใช้ หรือเข้าสู่ระบบด้วยบัญชีที่มีอยู่")
    add_step(flap, "3", "เริ่มทดลอง", "เลือกห้องทดลองและเริ่มเรียนรู้ได้ทันที")
    add_picture(flap, ASSETS / "scisiam-login-qr.png", width_cm=3.5, after=0)
    add_text(
        flap,
        "https://scisiam-app.vercel.app/login",
        size=12.5,
        bold=True,
        color=BLUE_DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=2,
    )
    add_note(flap, "ใช้งานผ่านเว็บ หรือเลือก “ติดตั้งแอป Windows” ที่หน้าเข้าสู่ระบบ")

    # Panel 6: back panel
    add_picture(back, ROOT / "public" / "scisiam-logo.png", width_cm=2.3, after=3)
    add_text(
        back,
        "เรียนรู้วิทยาศาสตร์\nผ่านการลงมือทำ",
        size=25,
        bold=True,
        color=NAVY,
        after=4,
        line=0.95,
    )
    add_text(
        back,
        "เปลี่ยนแนวคิดที่เข้าใจยากให้เป็นประสบการณ์ทดลองที่มองเห็น ปรับเปลี่ยน และเรียนรู้ได้ด้วยตนเอง",
        size=15.5,
        color=SLATE,
        after=5,
        line=1.08,
    )
    add_metrics(back)
    add_divider(back, after=4)
    add_check(back, "ข้อมูลและผลการทดลองเปลี่ยนแปลงแบบทันที")
    add_check(back, "บันทึกผลการทดลองและทบทวนประวัติการเรียนรู้")
    add_check(back, "ใช้งานได้ทั้งรายบุคคลและภายในชั้นเรียน")
    add_text(
        back,
        "ห้องปฏิบัติการเสมือนจริงภาษาไทย\nสำหรับการเรียนรู้ในยุคดิจิทัล",
        size=15,
        bold=True,
        color=BLUE_DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=5,
        line=1.0,
    )
    add_text(back, "© 2026 SciSiam Virtual Lab", size=11.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=4)

    # Panel 1: front cover
    add_picture(cover, ROOT / "public" / "scisiam-logo.png", width_cm=2.5, after=2)
    add_label(cover, "SCISIAM VIRTUAL LAB", fill=BLUE)
    add_text(
        cover,
        "ทดลองวิทยาศาสตร์\nได้ทุกที่ ทุกเวลา",
        size=28,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=5,
        after=2,
        line=0.9,
    )
    add_text(
        cover,
        "ห้องปฏิบัติการเสมือนจริง\nสำหรับนักเรียนและคุณครู",
        size=16,
        color=SLATE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=1,
        line=1.0,
    )
    add_picture(cover, ASSETS / "ai-oon-brochure.png", width_cm=6.4, after=1)
    add_label(cover, "เรียนรู้  •  ทดลอง  •  วิเคราะห์", fill=WHITE, color=BLUE_DARK)


def build_inside(document: Document) -> None:
    _, (intro, simulation, ai_classroom) = create_panel_table(document, (VERY_PALE, WHITE, VERY_PALE))

    # Panel 2: about SciSiam
    add_label(intro, "รู้จัก SCISIAM", fill=BLUE_DARK)
    add_text(intro, "SciSiam คืออะไร", size=25, bold=True, color=NAVY, before=5, after=4)
    add_text(
        intro,
        "SciSiam คือห้องปฏิบัติการวิทยาศาสตร์เสมือนจริงภาษาไทย ออกแบบเพื่อให้นักเรียนระดับมัธยมศึกษาได้ทดลองในสภาพแวดล้อมที่ปลอดภัย ควบคุมตัวแปร และสังเกตผลได้อย่างเป็นระบบ",
        size=15.2,
        color=SLATE,
        after=5,
        line=1.08,
    )
    add_text(intro, "ครอบคลุม 5 กลุ่มวิชา", size=18, bold=True, color=BLUE_DARK, after=3)
    add_subject(intro, "ฟิสิกส์", "Physics")
    add_subject(intro, "เคมี", "Chemistry")
    add_subject(intro, "ชีววิทยา", "Biology")
    add_subject(intro, "คณิตศาสตร์", "Mathematics")
    add_subject(intro, "พื้นฐานวิทยาศาสตร์", "Foundation")
    add_note(intro, "เรียนรู้จากการทดลองซ้ำได้ทุกเวลา โดยไม่จำกัดด้วยสถานที่หรืออุปกรณ์ในห้องปฏิบัติการ")

    # Panel 3: simulation experience
    add_label(simulation, "ประสบการณ์การทดลอง", fill=BLUE)
    add_text(simulation, "ทดลองจริง เห็นผลจริง", size=25, bold=True, color=NAVY, before=5, after=2)
    add_text(
        simulation,
        "ปรับตัวแปร ทดลองซ้ำ และสังเกตผลลัพธ์ได้ทันที พร้อมข้อมูลที่ช่วยเชื่อมโยงทฤษฎีกับสิ่งที่เกิดขึ้นจริง",
        size=15,
        color=SLATE,
        after=4,
        line=1.05,
    )
    add_picture(simulation, ASSETS / "simulation-preview.png", width_cm=8.35, after=1)
    add_text(
        simulation,
        "ตัวอย่าง: การจำลองกฎการเย็นตัวของนิวตัน",
        size=12.3,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=3,
    )
    add_check(simulation, "ภาพจำลองตอบสนองต่อค่าที่ผู้เรียนกำหนด")
    add_check(simulation, "แสดงค่าที่วัดได้ กราฟ และตารางอย่างเป็นระบบ")
    add_check(simulation, "บันทึกผลการทดลองเพื่อทบทวนภายหลัง")
    add_check(simulation, "เรียนรู้ความสัมพันธ์ระหว่างเหตุและผล")

    # Panel 4: AI and classroom
    add_label(ai_classroom, "เรียนรู้ร่วมกัน", fill=BLUE_DARK)
    add_text(ai_classroom, "AI ไออุ่นและ\nชั้นเรียนออนไลน์", size=25, bold=True, color=NAVY, before=5, after=4, line=0.92)
    add_info_card(
        ai_classroom,
        "AI ไออุ่น",
        "ผู้ช่วยระหว่างการเรียนรู้ ช่วยอธิบายแนวคิด ให้คำแนะนำระหว่างทดลอง และชวนผู้เรียนตรวจสอบคำตอบอย่างมีเหตุผล",
        image=ROOT / "public" / "ai-oon-logo.png",
    )
    add_text(ai_classroom, "", size=3, after=1)
    add_picture(ai_classroom, ASSETS / "classroom-preview.png", width_cm=8.25, after=1)
    add_text(
        ai_classroom,
        "คุณครูจัดห้องเรียน เลือกห้องทดลอง มอบหมายงาน และติดตามการส่งงานได้ในพื้นที่เดียว",
        size=14.2,
        color=SLATE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=3,
        line=1.03,
    )
    add_bullet(ai_classroom, "นักเรียน: ทดลอง บันทึกผล และทบทวนการเรียนรู้", size=14.5)
    add_bullet(ai_classroom, "คุณครู: จัดการชั้นเรียนและมอบหมายกิจกรรม", size=14.5)
    add_note(ai_classroom, "หมายเหตุ: AI ไออุ่นเป็นผู้ช่วยการเรียนรู้ ผู้ใช้ควรตรวจสอบข้อมูลร่วมกับบทเรียน แหล่งอ้างอิง หรือคุณครู")


def main() -> None:
    for path in (
        ASSETS / "ai-oon-brochure.png",
        ASSETS / "scisiam-login-qr.png",
        ASSETS / "simulation-preview.png",
        ASSETS / "classroom-preview.png",
        ROOT / "public" / "scisiam-logo.png",
        ROOT / "public" / "ai-oon-logo.png",
    ):
        if not path.exists():
            raise FileNotFoundError(path)

    document = Document()
    configure_document(document)
    build_outside(document)
    configure_section(document.add_section(WD_SECTION.NEW_PAGE))
    build_inside(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
