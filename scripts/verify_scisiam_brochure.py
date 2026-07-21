from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
from pypdf import PdfReader


ROOT = Path(r"D:\Scisiam_app")
DOCX = ROOT / "output" / "doc" / "แผ่นพับ_SciSiam_สำหรับกรรมการ.docx"
PDF = ROOT / "output" / "doc" / "แผ่นพับ_SciSiam_สำหรับกรรมการ.pdf"
ASSETS = ROOT / "output" / "doc" / "assets"
RENDERS = ROOT / "tmp" / "brochure-pages"


def check(condition: bool, message: str, failures: list[str]) -> None:
    if condition:
        print(f"PASS: {message}")
    else:
        failures.append(message)
        print(f"FAIL: {message}")


def main() -> None:
    failures: list[str] = []

    check(DOCX.exists() and DOCX.stat().st_size > 100_000, "DOCX exists and is non-empty", failures)
    check(PDF.exists() and PDF.stat().st_size > 100_000, "PDF exists and is non-empty", failures)

    with zipfile.ZipFile(DOCX) as archive:
        check(archive.testzip() is None, "DOCX package has no corrupt members", failures)
        document_xml = archive.read("word/document.xml").decode("utf-8")
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]

    document = Document(DOCX)
    check(len(document.tables) == 2, "DOCX has two brochure-side tables", failures)
    check(
        all(len(table.rows) == 1 and len(table.rows[0].cells) == 3 for table in document.tables),
        "Each brochure side contains three panels",
        failures,
    )
    check(len(document.sections) == 2, "DOCX has two print sections", failures)
    check(
        all(section.page_width > section.page_height for section in document.sections),
        "Both sections are landscape",
        failures,
    )
    check(len(media_files) >= 6, "DOCX embeds the required visual assets", failures)

    required_text = (
        "ทดลองวิทยาศาสตร์",
        "พร้อมทดลอง",
        "SciSiam คืออะไร",
        "ทดลองจริง เห็นผลจริง",
        "AI ไออุ่น",
        "ชั้นเรียนออนไลน์",
        "https://scisiam-app.vercel.app/login",
    )
    for text in required_text:
        check(text in document_xml, f"Required text is present: {text}", failures)

    check("ภารกิจนักวิทย์" not in document_xml, "Removed mission content is absent", failures)
    check("สิ้นสุดคู่มือ" not in document_xml, "Old manual ending text is absent", failures)

    text_runs = []
    for run in document.element.body.iter(qn("w:r")):
        text = "".join(node.text or "" for node in run.iter(qn("w:t")))
        if text.strip():
            text_runs.append(run)
    wrong_font_runs: list[str] = []
    for run in text_runs:
        text = "".join(node.text or "" for node in run.iter(qn("w:t"))).strip()
        r_pr = run.find(qn("w:rPr"))
        r_fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
        font_name = r_fonts.get(qn("w:ascii")) if r_fonts is not None else None
        if font_name != "TH Sarabun New":
            wrong_font_runs.append(text[:40])
    check(not wrong_font_runs, "Every visible DOCX text run uses TH Sarabun New", failures)

    pdf = PdfReader(PDF)
    check(len(pdf.pages) == 2, "PDF has exactly two pages", failures)
    for index, page in enumerate(pdf.pages, start=1):
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        check(width > height, f"PDF page {index} is landscape", failures)
        check(abs(width - 841.9) < 3 and abs(height - 595.3) < 3, f"PDF page {index} is A4 size", failures)

    font_names: set[str] = set()
    for page in pdf.pages:
        fonts = page["/Resources"].get("/Font") or {}
        for value in fonts.values():
            base_font = value.get_object().get("/BaseFont")
            if base_font:
                font_names.add(str(base_font))
    check(
        any("THSarabunNew" in name for name in font_names),
        "PDF embeds TH Sarabun New",
        failures,
    )

    for asset_name in (
        "ai-oon-brochure.png",
        "scisiam-login-qr.png",
        "simulation-preview.png",
        "classroom-preview.png",
    ):
        asset = ASSETS / asset_name
        check(asset.exists() and asset.stat().st_size > 1_000, f"Asset exists: {asset_name}", failures)
        if asset.exists():
            with Image.open(asset) as image:
                check(image.width >= 400 and image.height >= 400, f"Asset resolution is usable: {asset_name}", failures)

    render_files = sorted(RENDERS.glob("page-*.png"))
    check(len(render_files) == 2, "Visual verification produced two page renders", failures)
    for render in render_files:
        with Image.open(render) as image:
            check(image.width > image.height and image.width >= 1600, f"Rendered page is clear: {render.name}", failures)

    print(f"\nVerification summary: {len(failures)} failure(s)")
    if failures:
        for failure in failures:
            print(f"- {failure}")
        sys.exit(1)


if __name__ == "__main__":
    main()
